"""
Document Skill - 文档导入知识库

支持将文档内容解析并导入到知识库：
- TXT 纯文本
- PDF (需要安装 pypdf)
- Word/DOCX (需要安装 python-docx)
- HTML (需要安装 beautifulsoup4)
- Markdown
"""

import os
import re
from pathlib import Path

from assistant.skills.base import BaseSkill, ToolDefinition, register
from assistant.agent.rag import ingest_document


class DocumentSkill(BaseSkill):
    """文档处理技能 - 将文档内容导入知识库"""

    @property
    def name(self) -> str:
        return "document"

    @property
    def description(self) -> str:
        return "文档处理：解析TXT/PDF/Word/HTML/Markdown文件并导入知识库"

    def get_tools(self) -> list[ToolDefinition]:
        return [
            ToolDefinition(
                name="import_document",
                description="将文档内容导入知识库。支持txt/pdf/docx/html/md格式。",
                parameters={
                    "type": "object",
                    "properties": {
                        "file_path": {
                            "type": "string",
                            "description": "文档文件路径或URL",
                        },
                        "title": {
                            "type": "string",
                            "description": "知识库文档标题，如不提供则自动从文件名提取",
                        },
                    },
                    "required": ["file_path"],
                },
                handler=self._import_document,
                metadata={
                    "category": "write",
                    "side_effect": "data_write",
                    "required_all": ["file_path"],
                },
                keywords=["导入文档", "文档入库", "上传知识库"],
                intents=["import_document"],
            ),
            ToolDefinition(
                name="parse_document",
                description="解析文档内容并返回文本（不导入知识库）。用于预览文档内容。",
                parameters={
                    "type": "object",
                    "properties": {
                        "file_path": {
                            "type": "string",
                            "description": "文档文件路径或URL",
                        },
                        "max_lines": {
                            "type": "integer",
                            "description": "最大返回行数，默认100行",
                            "default": 100,
                        },
                    },
                    "required": ["file_path"],
                },
                handler=self._parse_document,
                metadata={
                    "category": "read",
                    "required_all": ["file_path"],
                },
                keywords=["解析文档", "预览文档", "读取PDF", "读取Word"],
                intents=["parse_document"],
            ),
        ]

    def _import_document(self, file_path: str, title: str = "") -> str:
        """将文档导入知识库"""
        try:
            # 解析文档内容
            content = self._extract_content(file_path)

            if not content or len(content.strip()) < 10:
                return f"文档内容为空或无法解析: {file_path}"

            # 生成标题
            if not title:
                title = Path(file_path).stem
                # 清理标题
                title = re.sub(r'[^\w\s\u4e00-\u9fa5]', '', title)
                title = title.strip()[:100]

            # 导入知识库
            result = ingest_document(
                title=title,
                content=content,
                source=file_path,
                doc_type=self._get_doc_type(file_path),
            )

            return (f"文档导入成功！\n"
                    f"- 标题: {title}\n"
                    f"- 分块数: {result.get('chunk_count', '?')}\n"
                    f"- 来源: {file_path}\n"
                    f"- 文档ID: {result.get('doc_id', '?')}")

        except Exception as e:
            return f"文档导入失败: {e}"

    def _parse_document(self, file_path: str, max_lines: int = 100) -> str:
        """解析文档内容（预览）"""
        try:
            content = self._extract_content(file_path)

            if not content:
                return "无法解析文档内容"

            lines = content.split('\n')
            if len(lines) > max_lines:
                content = '\n'.join(lines[:max_lines])
                content += f"\n... (还有 {len(lines) - max_lines} 行)"

            return f"文档内容预览:\n{content}"

        except Exception as e:
            return f"文档解析失败: {e}"

    def _extract_content(self, file_path: str) -> str:
        """根据文件类型提取内容"""
        # 处理 URL
        if file_path.startswith("http"):
            import requests
            response = requests.get(file_path, timeout=30)
            # 从 URL 推断文件类型
            content_type = response.headers.get("Content-Type", "")
            if "pdf" in content_type:
                return self._extract_pdf_bytes(response.content)
            elif "html" in content_type:
                return self._extract_html(response.text)
            else:
                return response.text

        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".txt" or suffix == ".text":
            return self._extract_txt(file_path)
        elif suffix == ".pdf":
            return self._extract_pdf(file_path)
        elif suffix in [".docx", ".doc"]:
            return self._extract_docx(file_path)
        elif suffix == ".html" or suffix == ".htm":
            return self._extract_html_file(file_path)
        elif suffix == ".md" or suffix == ".markdown":
            return self._extract_md(file_path)
        else:
            # 尝试作为纯文本读取
            return self._extract_txt(file_path)

    def _extract_txt(self, file_path: str) -> str:
        """读取 TXT 文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    def _extract_pdf(self, file_path: str) -> str:
        """读取 PDF 文件"""
        try:
            from pypdf import PdfReader
        except ImportError:
            return "需要安装 pypdf 库: pip install pypdf"

        reader = PdfReader(file_path)
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
        return "\n".join(texts)

    def _extract_pdf_bytes(self, data: bytes) -> str:
        """从 PDF 字节提取内容"""
        try:
            from pypdf import PdfReader
            from io import BytesIO
        except ImportError:
            return "需要安装 pypdf 库: pip install pypdf"

        reader = PdfReader(BytesIO(data))
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
        return "\n".join(texts)

    def _extract_docx(self, file_path: str) -> str:
        """读取 Word 文件"""
        try:
            from docx import Document
        except ImportError:
            return "需要安装 python-docx 库: pip install python-docx"

        doc = Document(file_path)
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)

        # 也读取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    texts.append(row_text)

        return "\n".join(texts)

    def _extract_html(self, html: str) -> str:
        """从 HTML 提取文本"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            # 简单处理：移除 HTML 标签
            import re
            text = re.sub(r'<[^>]+>', '', html)
            text = re.sub(r'\s+', ' ', text)
            return text.strip()

        soup = BeautifulSoup(html, "html.parser")

        # 移除脚本和样式
        for tag in soup(["script", "style"]):
            tag.decompose()

        # 获取文本
        text = soup.get_text(separator="\n")
        # 清理空白
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(line for line in lines if line)
        return text

    def _extract_html_file(self, file_path: str) -> str:
        """读取 HTML 文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            return self._extract_html(f.read())

    def _extract_md(self, file_path: str) -> str:
        """读取 Markdown 文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        # 可选：移除 Markdown 语法（简单处理）
        import re
        # 移除图片
        content = re.sub(r'!\[.*?\]\(.*?\)', '', content)
        # 移除链接
        content = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', content)
        # 移除标题标记但保留文字
        content = re.sub(r'^#+ ', '', content, flags=re.MULTILINE)
        # 移除加粗斜体标记
        content = re.sub(r'[*_]{1,3}(.*?)[*_]{1,3}', r'\1', content)
        # 移除代码块标记
        content = re.sub(r'```.*?\n', '', content)

        return content.strip()

    def _get_doc_type(self, file_path: str) -> str:
        """获取文档类型"""
        suffix = Path(file_path).suffix.lower()
        type_map = {
            ".txt": "text",
            ".text": "text",
            ".pdf": "pdf",
            ".docx": "docx",
            ".doc": "doc",
            ".html": "html",
            ".htm": "html",
            ".md": "markdown",
            ".markdown": "markdown",
        }
        return type_map.get(suffix, "text")


register(DocumentSkill)
